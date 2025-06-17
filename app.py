import streamlit as st
from utils.transcript_utils import extract_video_id, get_spanish_transcript
import io

st.set_page_config(page_title="YouTube Spanish Transcript", layout="centered")
st.title("ðŸ“¹ YouTube Spanish Transcript Extractor")

url = st.text_input("Enter YouTube video URL:")
custom_title = st.text_input("Enter a name for the transcript file (without extension):", value="transcript")

if url:
    video_id = extract_video_id(url)

    if video_id:
        with st.spinner("Fetching transcript..."):
            transcript, error = get_spanish_transcript(video_id)

        if error:
            st.error(error)
        else:
            st.success("Transcript retrieved successfully!")

            full_text = "\n".join([entry['text'] for entry in transcript])
            st.text_area("ðŸ“‹ Transcript", full_text, height=300)

            file_name = f"{custom_title.strip() or 'transcript'}.txt"
            st.download_button(
                label="ðŸ“¥ Download Transcript as .txt",
                data=full_text,
                file_name=file_name,
                mime="text/plain"
            )

            try:
                from docx import Document
                doc = Document()
                doc.add_heading(custom_title, 0)
                for line in full_text.split('\n'):
                    doc.add_paragraph(line)
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="ðŸ“„ Download as .docx",
                    data=buffer,
                    file_name=f"{custom_title.strip() or 'transcript'}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except ImportError:
                st.info("Install `python-docx` to enable .docx downloads: `pip install python-docx`")
    else:
        st.error("Invalid YouTube URL.")
